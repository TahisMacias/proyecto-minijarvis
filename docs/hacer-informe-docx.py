"""Convierte docs/informe-tecnico.md en el Word que se entrega.

    python docs/hacer-informe-docx.py docs/informe-tecnico.md "docs/Informe tecnico.docx"

El enunciado pide un informe de 4 a 8 PAGINAS. Un .md no tiene paginas: no se puede
entregar asi ni se puede contar. Esto produce el documento que se entrega.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ENTRADA = Path(sys.argv[1])
SALIDA = Path(sys.argv[2])

TURQUESA = RGBColor(0x0E, 0x7C, 0x74)
GRIS = RGBColor(0x44, 0x44, 0x4A)

doc = Document()

# Pagina y tipografia base. Letter, cuerpo en 10.5 pt. Con 11 pt y margenes de 2.5 cm
# el informe salia en 9 paginas y el enunciado pide de 4 a 8, asi que se aprieta el
# formato en vez de recortar contenido: se pierde aire, no argumentos.
seccion = doc.sections[0]
seccion.page_width = Cm(21.59)
seccion.page_height = Cm(27.94)
for lado in ("left_margin", "right_margin"):
    setattr(seccion, lado, Cm(2.2))
seccion.top_margin = Cm(2.0)
seccion.bottom_margin = Cm(2.0)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08

for nivel, tam in ((1, 17), (2, 13), (3, 11.5)):
    estilo = doc.styles[f"Heading {nivel}"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(tam)
    estilo.font.color.rgb = TURQUESA
    estilo.font.bold = True
    estilo.paragraph_format.space_before = Pt(12 if nivel < 3 else 8)
    estilo.paragraph_format.space_after = Pt(6)


def negritas_y_codigo(parrafo, texto):
    """Convierte **negrita** y `codigo` de markdown en formato real de Word."""
    for trozo in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", texto):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            parrafo.add_run(trozo[2:-2]).bold = True
        elif trozo.startswith("`") and trozo.endswith("`"):
            r = parrafo.add_run(trozo[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8B, 0x2E, 0x5C)
        else:
            parrafo.add_run(trozo)


lineas = ENTRADA.read_text(encoding="utf-8").split("\n")
i = 0
primer_titulo = True

while i < len(lineas):
    linea = lineas[i]

    # --- bloque de codigo: se respeta tal cual, en monoespaciada ---
    if linea.startswith("```"):
        i += 1
        cuerpo = []
        while i < len(lineas) and not lineas[i].startswith("```"):
            cuerpo.append(lineas[i])
            i += 1
        i += 1
        for fila in cuerpo:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(fila if fila.strip() else " ")
            r.font.name = "Consolas"
            # El diagrama de la arquitectura es ancho: a 7 pt cabe sin partirse.
            r.font.size = Pt(7)
            r.font.color.rgb = GRIS
        doc.add_paragraph()
        continue

    # --- tabla ---
    if linea.startswith("|") and i + 1 < len(lineas) and set(lineas[i + 1].replace("|", "").strip()) <= set("-: "):
        filas = []
        while i < len(lineas) and lineas[i].startswith("|"):
            celdas = [c.strip() for c in lineas[i].strip().strip("|").split("|")]
            if not set("".join(celdas)) <= set("-: "):
                filas.append(celdas)
            i += 1
        tabla = doc.add_table(rows=0, cols=len(filas[0]))
        tabla.style = "Light Grid Accent 1"
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        for numero, fila in enumerate(filas):
            celdas_word = tabla.add_row().cells
            for celda, texto in zip(celdas_word, fila):
                celda.text = ""
                p = celda.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                negritas_y_codigo(p, texto)
                for r in p.runs:
                    r.font.size = Pt(9)
                    if numero == 0:
                        r.bold = True
        doc.add_paragraph()
        continue

    # --- separador horizontal: se ignora, los titulos ya separan ---
    if linea.strip() == "---":
        i += 1
        continue

    # --- titulos ---
    if linea.startswith("### "):
        doc.add_heading(linea[4:].strip(), level=3)
    elif linea.startswith("## "):
        doc.add_heading(linea[3:].strip(), level=2)
    elif linea.startswith("# "):
        if primer_titulo:
            t = doc.add_heading(linea[2:].strip(), level=1)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            primer_titulo = False
        else:
            doc.add_heading(linea[2:].strip(), level=1)

    # --- vinetas ---
    elif linea.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        negritas_y_codigo(p, linea[2:].strip())

    # --- parrafo: se juntan las lineas hasta el proximo salto en blanco ---
    elif linea.strip():
        trozos = [linea.strip()]
        while (i + 1 < len(lineas) and lineas[i + 1].strip()
               and not lineas[i + 1].startswith(("#", "-", "|", "```"))):
            i += 1
            trozos.append(lineas[i].strip())
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        negritas_y_codigo(p, " ".join(trozos))

    i += 1

doc.save(SALIDA)
print(f"escrito: {SALIDA}")
print(f"tamano : {SALIDA.stat().st_size // 1024} KB")
